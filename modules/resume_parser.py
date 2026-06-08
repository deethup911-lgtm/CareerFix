import os
import fitz  # PyMuPDF
import docx

def parse_pdf(file_path_or_bytes):
    try:
        text = ""
        if isinstance(file_path_or_bytes, bytes):
            doc = fitz.open(stream=file_path_or_bytes, filetype="pdf")
        else:
            doc = fitz.open(file_path_or_bytes)
        for page in doc:
            text += page.get_text("text") + "\n"
        return text
    except Exception as e:
        print(f"Error parsing PDF: {e}")
        return ""

def parse_docx(file_path_or_bytes):
    try:
        if isinstance(file_path_or_bytes, bytes):
            from io import BytesIO
            doc = docx.Document(BytesIO(file_path_or_bytes))
        else:
            doc = docx.Document(file_path_or_bytes)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return ""

def extract_text(file_obj, filename):
    text = ""
    try:
        file_bytes = file_obj.read()
        if filename.lower().endswith('.pdf'):
            text = parse_pdf(file_bytes)
        elif filename.lower().endswith('.docx'):
            text = parse_docx(file_bytes)
        else:
            print("Unsupported file format.")
    except Exception as e:
        print(f"Error extracting text: {e}")
    return text.strip()
